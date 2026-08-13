from __future__ import annotations

import copy
import html
import json
import os
import shutil
import subprocess
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from openpyxl import load_workbook
from PIL import Image

from .accounting import (
    VALIDATION_RULE_VERSION,
    Evidence,
    Transaction,
    processing_method_requires_evidence,
    public_monthly_summary,
    validate_ledger,
)
from .integrity import canonical_sha256, file_sha256
from .media import image_dimensions, render_media_pages
from .naming import normalize_filename, normalize_nfc, normalize_relative_path


TEMPLATE_VERSION = "atlas-documents-v1.4-combined-official-template-fill"


def _won(value: int) -> str:
    return f"{value:,}원"


def _account_label(account_id: str | None) -> str:
    labels = {
        "primary": "동아리운영계좌(우리은행)",
        "dues_intake": "회비입금계좌(IBK기업은행)",
    }
    return labels.get(account_id or "primary", account_id or "primary")


def _effective_row_capacity(requested: int, transaction_count: int) -> int:
    required = max(requested, transaction_count)
    for capacity in (40, 80, 120):
        if required <= capacity:
            return capacity
    raise ValueError("공식 수입지출관리대장 양식은 최대 120칸까지 지원합니다.")


def _normalized(value: str) -> str:
    return normalize_nfc(value)


def _official_template_dir() -> Path:
    candidates = []
    if os.getenv("ATLAS_TEMPLATE_ROOT"):
        candidates.append(Path(os.environ["ATLAS_TEMPLATE_ROOT"]))
    candidates.extend((Path("/app/sample"), Path(__file__).resolve().parents[3] / "sample"))
    for candidate in candidates:
        if candidate.is_dir() and "표준회계양식" in _normalized(candidate.name):
            return candidate
        if not candidate.is_dir():
            continue
        for child in candidate.iterdir():
            if child.is_dir() and "표준회계양식" in _normalized(child.name):
                return child
    raise FileNotFoundError("동아리연합회 공식 회계양식 디렉터리를 찾을 수 없습니다.")


def _official_template(kind: str, capacity: int | None = None) -> Path:
    for path in _official_template_dir().iterdir():
        name = _normalized(path.name)
        if kind == "ledger" and path.suffix.lower() == ".xls" and f"({capacity}칸)" in name:
            return path
        if kind == "evidence" and path.suffix.lower() == ".docx" and "영수증 및 소명자료" in name:
            return path
        if kind == "account" and path.suffix.lower() == ".docx" and "계좌전체내역" in name:
            return path
    raise FileNotFoundError(f"동아리연합회 공식 양식을 찾을 수 없습니다: {kind}")


def _convert_official_ledger_template(template_path: Path, destination: Path) -> None:
    soffice = os.getenv("ATLAS_SOFFICE_BIN") or shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("공식 .xls 양식 변환에 필요한 LibreOffice를 찾을 수 없습니다.")
    destination.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="atlas-ledger-template-") as temporary:
        workdir = Path(temporary)
        profile = workdir / "profile"
        converted_dir = workdir / "converted"
        profile.mkdir()
        converted_dir.mkdir()
        process = subprocess.run(
            [
                soffice,
                f"-env:UserInstallation={profile.resolve().as_uri()}",
                "--headless",
                "--convert-to",
                "xlsx",
                "--outdir",
                str(converted_dir),
                str(template_path),
            ],
            capture_output=True,
            text=True,
            timeout=90,
            check=False,
        )
        converted = next(converted_dir.glob("*.xlsx"), None)
        if process.returncode or converted is None:
            detail = (process.stderr or process.stdout or "unknown conversion error").strip()
            raise RuntimeError(f"공식 수입지출관리대장 양식을 변환할 수 없습니다: {detail}")
        shutil.copy2(converted, destination)


def _populate_ledger_sheet(
    ws,
    club_name: str,
    semester: str,
    treasurer_name: str,
    president_name: str,
    opening_balance: int,
    transactions: list[Transaction],
    row_capacity: int,
    account_label: str | None = None,
) -> None:
    title_suffix = f" ({account_label})" if account_label else ""
    ws["B2"] = f"{club_name} 수입지출 관리대장{title_suffix}"
    ws["R3"] = f"{treasurer_name}  (인)"
    ws["V3"] = f"{president_name}  (인)"
    semester_label = next((label for label in ("1학기", "2학기") if label in semester), semester)
    ws["B7"] = semester_label
    ws["T7"] = opening_balance

    start_row = 10
    for index in range(row_capacity):
        row = start_row + index
        if index < len(transactions):
            tx = transactions[index]
            try:
                transaction_date = datetime.fromisoformat(tx.date).replace(tzinfo=None)
            except ValueError:
                transaction_date = tx.date
            ws.cell(row, 2, tx.number)
            ws.cell(row, 4, transaction_date)
            ws.cell(row, 7, tx.description)
            ws.cell(row, 11, tx.income or None)
            ws.cell(row, 17, tx.expense or None)
            ws.cell(row, 24, tx.note)
        else:
            ws.cell(row, 2, index + 1)
            for column in (4, 7, 11, 17, 24):
                ws.cell(row, column, None)


def generate_ledger_workbook(
    output_path: Path,
    club_name: str,
    semester: str,
    treasurer_name: str,
    president_name: str,
    opening_balance: int,
    transactions: list[Transaction],
    row_capacity: int,
) -> dict:
    row_capacity = _effective_row_capacity(row_capacity, len(transactions))
    template_path = _official_template("ledger", row_capacity)
    _convert_official_ledger_template(template_path, output_path)
    wb = load_workbook(output_path)
    _populate_ledger_sheet(
        wb.active,
        club_name,
        semester,
        treasurer_name,
        president_name,
        opening_balance,
        transactions,
        row_capacity,
    )
    if hasattr(wb, "calculation"):
        wb.calculation.fullCalcOnLoad = True
        wb.calculation.forceFullCalc = True
    wb.save(output_path)
    return {"template_filename": _normalized(template_path.name), "template_sha256": file_sha256(template_path)}


def generate_combined_ledger_workbook(
    output_path: Path,
    club_name: str,
    semester: str,
    treasurer_name: str,
    president_name: str,
    accounts: dict[str, dict],
    row_capacity: int,
) -> dict:
    effective_capacity = _effective_row_capacity(
        row_capacity,
        max((len(bundle["transactions"]) for bundle in accounts.values()), default=0),
    )
    template_path = _official_template("ledger", effective_capacity)
    _convert_official_ledger_template(template_path, output_path)
    wb = load_workbook(output_path)
    primary_sheet = wb.active
    dues_sheet = wb.copy_worksheet(primary_sheet)
    primary_sheet.title = "동아리운영계좌(우리은행)"
    dues_sheet.title = "회비입금계좌(IBK기업은행)"
    sheets = {"primary": primary_sheet, "dues_intake": dues_sheet}
    sheet_labels = {"primary": _account_label("primary"), "dues_intake": _account_label("dues_intake")}
    coverage: dict[str, dict] = {}
    for account_id, ws in sheets.items():
        bundle = accounts[account_id]
        _populate_ledger_sheet(
            ws,
            club_name,
            semester,
            treasurer_name,
            president_name,
            int(bundle.get("opening_balance") or 0),
            bundle["transactions"],
            effective_capacity,
            sheet_labels[account_id],
        )
        coverage[account_id] = {
            "sheet_name": ws.title,
            "transaction_rows": len(bundle["transactions"]),
            "row_capacity": effective_capacity,
        }
    if hasattr(wb, "calculation"):
        wb.calculation.fullCalcOnLoad = True
        wb.calculation.forceFullCalc = True
    wb.save(output_path)
    return {
        "template_filename": _normalized(template_path.name),
        "template_sha256": file_sha256(template_path),
        "accounts": coverage,
    }


def _add_contained_picture(paragraph, image_path: Path, max_width: float, max_height: float) -> None:
    width, height = image_dimensions(image_path)
    scale = min(max_width / width, max_height / height)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(
        str(image_path),
        width=Inches(width * scale),
        height=Inches(height * scale),
    )


def _linked_evidence(tx: Transaction, evidence: list[Evidence]) -> list[Evidence]:
    linked: list[Evidence] = []
    ids = set(tx.evidence_ids)
    if tx.evidence_id:
        ids.add(tx.evidence_id)
    for item in evidence:
        if item.account_id == tx.account_id and (
            item.id in ids or item.transaction_number == tx.number or tx.transaction_id in item.transaction_ids
        ):
            if item not in linked:
                linked.append(item)
    return linked


def _set_paragraph_text(paragraph, value: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def _set_cell_text(cell, value: str) -> None:
    _set_paragraph_text(cell.paragraphs[0], value)
    for paragraph in cell.paragraphs[1:]:
        _set_paragraph_text(paragraph, "")


def _fill_image_cell(cell, image_path: Path, label: str, max_width: float, max_height: float) -> None:
    _set_cell_text(cell, label)
    for extra_paragraph in list(cell.paragraphs[1:]):
        extra_paragraph._element.getparent().remove(extra_paragraph._element)
    paragraph = cell.paragraphs[0]
    if label:
        paragraph.add_run("\n")
    _add_contained_picture(paragraph, image_path, max_width=max_width, max_height=max_height)


def _append_evidence_template_pages(doc: Document, required_tables: int) -> None:
    while len(doc.tables) < required_tables:
        table_element = doc.tables[-1]._tbl
        before = table_element.getprevious()
        after = table_element.getnext()
        section = doc.element.body.sectPr
        for element in (before, table_element, after):
            section.addprevious(copy.deepcopy(element))


def _append_account_template_pages(doc: Document, required_tables: int) -> None:
    while len(doc.tables) < required_tables:
        table_element = doc.tables[-1]._tbl
        first_signature = table_element.getnext()
        second_signature = first_signature.getnext()
        section = doc.element.body.sectPr
        for element in (table_element, first_signature, second_signature):
            section.addprevious(copy.deepcopy(element))


def _populate_evidence_slot(table, slot_index: int, entry: dict, treasurer_name: str, reviewer_name: str | None) -> None:
    base = 0 if slot_index == 0 else 5
    tx: Transaction | None = entry.get("transaction")
    item: Evidence | None = entry.get("evidence")
    image_path: Path | None = entry.get("page")
    date_value = tx.date if tx else (item.evidence_date if item else None)
    number_value = tx.number if tx else (item.transaction_number if item else None)
    description = tx.description if tx else (item.filename if item else "미연결 자료")
    note = tx.note if tx and tx.note else (entry.get("message") or (item.filename if item else ""))
    amount = (tx.expense or tx.income) if tx else (item.amount if item else None)
    if tx and "".join(tx.processing_method.split()) in {"카드결제취소", "이자입금취소"}:
        amount = -abs(amount or 0)

    _set_cell_text(table.cell(base, 1), f"날짜 : {date_value or '-'}")
    _set_cell_text(table.cell(base, 3), f"번호(수입지출관리대장) : {number_value or '-'}")
    _set_cell_text(table.cell(base + 1, 1), f"지출자 : {treasurer_name}  (인)")
    _set_cell_text(table.cell(base + 1, 3), f"검토자 : {reviewer_name or '-'}  (인)")
    _set_cell_text(table.cell(base + 2, 1), f"사업명 : {description}")
    _set_cell_text(table.cell(base + 3, 1), f"내용 & 특이사항\n{note or '특이사항 없음'}")
    _set_cell_text(table.cell(base + 4, 1), "지출 금액")
    _set_cell_text(table.cell(base + 4, 2), _won(int(amount or 0)))
    image_cell = table.cell(base, 0)
    if image_path:
        _fill_image_cell(image_cell, image_path, "", max_width=3.45, max_height=3.55)
    else:
        _set_cell_text(image_cell, entry.get("message") or "연결된 영수증 또는 소명자료가 없습니다.")


def generate_evidence_document(
    output_path: Path,
    club_name: str,
    treasurer_name: str,
    reviewer_name: str | None,
    transactions: list[Transaction],
    evidence: list[Evidence],
) -> dict:
    template_path = _official_template("evidence")
    doc = Document(template_path)
    render_dir = output_path.parent / ".rendered-evidence"
    expense_rows = [item for item in transactions if item.expense]
    used_evidence: set[str] = set()
    embedded_files = 0
    embedded_pages = 0
    unsupported_files: list[str] = []
    missing_required_cards: list[int] = []
    account_breakdown: dict[str, dict[str, int | str]] = {}
    entries: list[dict] = []

    def count_account(item: Evidence, pages: int = 0) -> None:
        row = account_breakdown.setdefault(
            item.account_id,
            {"account_id": item.account_id, "label": _account_label(item.account_id), "files": 0, "pages": 0},
        )
        row["files"] = int(row["files"]) + 1
        row["pages"] = int(row["pages"]) + pages

    for tx in transactions:
        linked = [item for item in _linked_evidence(tx, evidence) if item.kind != "account_capture"]
        if not linked:
            if processing_method_requires_evidence(tx.processing_method):
                missing_required_cards.append(tx.number)
                entries.append(
                    {
                        "transaction": tx,
                        "message": f"장부 {tx.number}번 {tx.processing_method.strip()} 거래에 매칭되는 소명자료 또는 영수증이 존재하지 않습니다.",
                    }
                )
            continue
        for item in linked:
            used_evidence.add(item.id)
            source = Path(item.local_path) if item.local_path else None
            pages = render_media_pages(source, render_dir / item.id) if source and source.exists() else []
            if not pages:
                unsupported_files.append(item.filename)
                count_account(item)
                entries.append({"transaction": tx, "evidence": item, "message": f"{item.filename}\n원본은 제출 ZIP에 포함됩니다."})
                continue
            embedded_files += 1
            count_account(item, len(pages))
            for page in pages:
                entries.append({"transaction": tx, "evidence": item, "page": page})
                embedded_pages += 1

    unlinked = [item for item in evidence if item.kind != "account_capture" and item.id not in used_evidence]
    for item in unlinked:
        source = Path(item.local_path) if item.local_path else None
        pages = render_media_pages(source, render_dir / item.id) if source and source.exists() else []
        if not pages:
            unsupported_files.append(item.filename)
            count_account(item)
            entries.append({"evidence": item, "message": f"미연결 자료: {item.filename}\n원본은 제출 ZIP에 포함됩니다."})
            continue
        embedded_files += 1
        count_account(item, len(pages))
        for page in pages:
            entries.append({"evidence": item, "page": page, "message": f"미연결 자료: {item.filename}"})
            embedded_pages += 1

    required_tables = max(1, (len(entries) + 1) // 2)
    _append_evidence_template_pages(doc, required_tables)
    for index, entry in enumerate(entries):
        _populate_evidence_slot(doc.tables[index // 2], index % 2, entry, treasurer_name, reviewer_name)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    shutil.rmtree(render_dir, ignore_errors=True)
    return {
        "expense_count": len(expense_rows),
        "missing_expense_numbers": [number for number in missing_required_cards if any(tx.number == number and tx.expense for tx in transactions)],
        "missing_required_card_numbers": missing_required_cards,
        "embedded_files": embedded_files,
        "embedded_pages": embedded_pages,
        "unsupported_files": sorted(set(unsupported_files)),
        "account_breakdown": list(account_breakdown.values()),
        "template_filename": _normalized(template_path.name),
        "template_sha256": file_sha256(template_path),
    }


def generate_account_capture_document(
    output_path: Path,
    club_name: str,
    president_name: str,
    evidence: list[Evidence],
) -> dict:
    template_path = _official_template("account")
    doc = Document(template_path)
    captures = [item for item in evidence if item.kind == "account_capture"]
    render_dir = output_path.parent / ".rendered-account-captures"
    embedded_pages = 0
    unsupported_files: list[str] = []
    account_breakdown: dict[str, dict[str, int | str]] = {}
    entries: list[dict] = []
    for item in captures:
        breakdown = account_breakdown.setdefault(item.account_id, {"account_id": item.account_id, "label": _account_label(item.account_id), "files": 0, "pages": 0})
        source = Path(item.local_path) if item.local_path else None
        pages = render_media_pages(source, render_dir / item.id) if source and source.exists() else []
        breakdown["files"] = int(breakdown["files"]) + 1
        if not pages:
            unsupported_files.append(item.filename)
            entries.append({"evidence": item, "message": f"{item.filename}\n원본은 제출 ZIP에 포함됩니다."})
            continue
        breakdown["pages"] = int(breakdown["pages"]) + len(pages)
        for page in pages:
            entries.append({"evidence": item, "page": page})
            embedded_pages += 1

    required_tables = max(1, (len(entries) + 3) // 4)
    _append_account_template_pages(doc, required_tables)
    for paragraph in doc.paragraphs:
        normalized_text = _normalized(paragraph.text)
        if "동아리" in normalized_text:
            _set_paragraph_text(paragraph, f"{club_name} ")
        elif "회장" in normalized_text:
            _set_paragraph_text(paragraph, f"회장 {president_name}  (인) ")
    for index, entry in enumerate(entries):
        table = doc.tables[index // 4]
        position = index % 4
        cell = table.cell(position // 2, position % 2)
        item = entry["evidence"]
        label = f"<{index + 1}> {_account_label(item.account_id)}"
        if entry.get("page"):
            _fill_image_cell(cell, entry["page"], label, max_width=3.0, max_height=3.55)
        else:
            _set_cell_text(cell, f"{label}\n{entry['message']}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    shutil.rmtree(render_dir, ignore_errors=True)
    return {
        "capture_file_count": len(captures),
        "embedded_capture_pages": embedded_pages,
        "embedded_total_pages": embedded_pages,
        "unsupported_files": sorted(set(unsupported_files)),
        "account_breakdown": list(account_breakdown.values()),
        "template_filename": _normalized(template_path.name),
        "template_sha256": file_sha256(template_path),
    }


def _validation_issue_html(issue: dict) -> str:
    ledger_numbers = issue.get("transaction_numbers") or []
    if issue.get("transaction_number") is not None:
        ledger_numbers = [issue["transaction_number"]]
    ledger_id = ""
    if ledger_numbers:
        label = ", ".join(str(int(number)) for number in ledger_numbers)
        ledger_id = f'<span class="ledger-id">장부 ID: {html.escape(label)}</span> '
    return (
        f"<li><strong>{html.escape(issue['severity'])}</strong> {ledger_id}"
        f"{html.escape(issue['code'])}: {html.escape(issue['message'])}</li>"
    )


def generate_validation_report(output_path: Path, validation: dict) -> None:
    summary = validation["summary"]
    rows = "\n".join(_validation_issue_html(issue) for issue in validation["issues"])
    output_path.write_text(
        f"""<!doctype html>
<html lang="ko">
<head>
  <meta charset="utf-8" />
  <title>ATLAS 검증 리포트</title>
  <style>
    body {{ font-family: Arial, sans-serif; margin: 40px; color: #172026; }}
    h1 {{ margin-bottom: 8px; }}
    table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
    th, td {{ border: 1px solid #c7d0d9; padding: 10px; text-align: left; }}
    th {{ background: #e9f3df; }}
    .status {{ font-weight: 700; }}
    .ledger-id {{ display: inline-block; margin: 0 6px; padding: 2px 6px; background: #eef2f1; font-weight: 700; }}
  </style>
</head>
<body>
  <h1>ATLAS 검증 리포트</h1>
  <p class="status">상태: {html.escape(validation['status'])}</p>
  <table>
    <tr><th>이전잔액</th><td>{summary['opening_balance']:,}원</td></tr>
    <tr><th>수입총액</th><td>{summary['total_income']:,}원</td></tr>
    <tr><th>지출총액</th><td>{summary['total_expense']:,}원</td></tr>
    <tr><th>계산 최종잔액</th><td>{summary['computed_closing_balance']:,}원</td></tr>
    <tr><th>장부 최종잔액</th><td>{summary['reported_closing_balance']:,}원</td></tr>
  </table>
  <h2>검증 항목</h2>
  <ul>{rows or '<li>검증 이슈가 없습니다.</li>'}</ul>
</body>
</html>
""",
        encoding="utf-8",
    )


def build_submission_package(
    output_root: Path,
    package_id: str,
    club_name: str,
    semester: str,
    treasurer_name: str,
    president_name: str,
    reviewer_name: str | None,
    opening_balance: int,
    expected_closing_balance: int | None,
    transactions: list[Transaction],
    evidence: list[Evidence],
    row_capacity: int,
    created_by: str = "system",
    snapshot_id: str | None = None,
    ledger_data_hash: str | None = None,
    period_start: str | None = None,
    period_end: str | None = None,
    source_files: list[dict] | None = None,
) -> dict:
    package_dir = output_root / package_id
    package_dir.mkdir(parents=True, exist_ok=True)
    validation = validate_ledger(
        opening_balance,
        transactions,
        evidence,
        expected_closing_balance,
        period_start,
        period_end,
    )

    ledger_path = package_dir / "수입지출관리대장.xlsx"
    evidence_path = package_dir / "영수증_및_소명자료.docx"
    captures_path = package_dir / "계좌전체내역.docx"
    report_path = package_dir / "검증_리포트.html"
    manifest_path = package_dir / "manifest.json"
    zip_path = output_root / f"{package_id}.zip"

    effective_capacity = _effective_row_capacity(row_capacity, len(transactions))
    ledger_coverage = generate_ledger_workbook(
        ledger_path,
        club_name,
        semester,
        treasurer_name,
        president_name,
        opening_balance,
        transactions,
        effective_capacity,
    )
    evidence_coverage = generate_evidence_document(
        evidence_path,
        club_name,
        treasurer_name,
        reviewer_name,
        transactions,
        evidence,
    )
    account_coverage = generate_account_capture_document(
        captures_path,
        club_name,
        president_name,
        evidence,
    )
    generate_validation_report(report_path, validation)

    copied_evidence: list[dict] = []
    evidence_dir = package_dir / "증빙자료"
    evidence_dir.mkdir(exist_ok=True)
    for item in evidence:
        if item.local_path and Path(item.local_path).exists():
            source_path = Path(item.local_path)
            target = evidence_dir / f"{item.id}_{normalize_filename(item.filename or source_path.name)}"
            shutil.copy2(item.local_path, target)
            copied_evidence.append(
                {
                    "id": item.id,
                    "filename": normalize_filename(target.name),
                    "relative_path": normalize_relative_path(target.relative_to(package_dir)),
                    "sha256": file_sha256(target),
                    "transaction_ids": list(item.transaction_ids),
                    "transaction_number": item.transaction_number,
                }
            )

    copied_sources: list[dict] = []
    source_dir = package_dir / "원본자료"
    source_dir.mkdir(exist_ok=True)
    for item in source_files or []:
        if item.get("kind") != "ledger_workbook":
            continue
        source_path = Path(str(item.get("local_path") or ""))
        if not source_path.exists():
            continue
        source_filename = normalize_filename(str(item.get("filename") or source_path.name))
        target = source_dir / f"{item.get('kind', 'source')}_{source_filename}"
        shutil.copy2(source_path, target)
        copied_sources.append(
            {
                "kind": item.get("kind", "source"),
                "filename": normalize_filename(target.name),
                "relative_path": normalize_relative_path(target.relative_to(package_dir)),
                "sha256": file_sha256(target),
            }
        )

    generated_paths = [ledger_path, evidence_path, captures_path, report_path]
    file_entries = [
        {
            "relative_path": normalize_relative_path(path.relative_to(package_dir)),
            "size": path.stat().st_size,
            "sha256": file_sha256(path),
        }
        for path in generated_paths
    ]
    file_entries.extend(
        {
            "relative_path": item["relative_path"],
            "size": (package_dir / item["relative_path"]).stat().st_size,
            "sha256": item["sha256"],
        }
        for item in copied_evidence
    )
    file_entries.extend(
        {
            "relative_path": item["relative_path"],
            "size": (package_dir / item["relative_path"]).stat().st_size,
            "sha256": item["sha256"],
        }
        for item in copied_sources
    )

    document_coverage = {
        "ledger_transaction_rows": len(transactions),
        "ledger_row_capacity": effective_capacity,
        "ledger_template": ledger_coverage,
        "evidence_document": evidence_coverage,
        "account_document": account_coverage,
        "copied_evidence_files": len(copied_evidence),
        "copied_source_files": len(copied_sources),
    }

    manifest = {
        "package_id": package_id,
        "club_name": club_name,
        "semester": semester,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "generated_by": created_by,
        "snapshot_id": snapshot_id,
        "ledger_data_hash": ledger_data_hash,
        "validation_rule_version": VALIDATION_RULE_VERSION,
        "template_version": TEMPLATE_VERSION,
        "document_coverage": document_coverage,
        "validation": validation,
        "evidence": copied_evidence,
        "source_files": copied_sources,
        "files": file_entries,
    }
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in package_dir.rglob("*"):
            if path.is_file():
                archive.write(path, normalize_relative_path(path.relative_to(package_dir)))

    zip_hash = file_sha256(zip_path)
    manifest_hash = file_sha256(manifest_path)

    return {
        "zip_path": str(zip_path),
        "zip_sha256": zip_hash,
        "manifest_sha256": manifest_hash,
        "manifest": manifest,
        "document_coverage": document_coverage,
        "row_capacity": effective_capacity,
        "validation": validation,
        "artifacts": [
            {"kind": "ledger", "path": str(ledger_path), "sha256": file_sha256(ledger_path)},
            {"kind": "evidence", "path": str(evidence_path), "sha256": file_sha256(evidence_path)},
            {"kind": "account_captures", "path": str(captures_path), "sha256": file_sha256(captures_path)},
            {"kind": "validation_report", "path": str(report_path), "sha256": file_sha256(report_path)},
            {"kind": "manifest", "path": str(manifest_path), "sha256": manifest_hash},
            {"kind": "zip", "path": str(zip_path), "sha256": zip_hash},
        ],
    }


def _combined_validation(accounts: dict[str, dict]) -> dict:
    account_results: dict[str, dict] = {}
    issues: list[dict] = []
    for account_id, bundle in accounts.items():
        result = validate_ledger(
            int(bundle.get("opening_balance") or 0),
            bundle["transactions"],
            bundle["evidence"],
            bundle.get("expected_closing_balance"),
            bundle.get("period_start"),
            bundle.get("period_end"),
        )
        account_results[account_id] = result
        for issue in result["issues"]:
            issues.append(
                {
                    **issue,
                    "code": f"{account_id}:{issue['code']}",
                    "account_id": account_id,
                    "account_label": _account_label(account_id),
                }
            )
    error_count = sum(1 for issue in issues if issue["severity"] == "ERROR" and issue.get("status") != "ACKNOWLEDGED")
    warning_count = sum(1 for issue in issues if issue["severity"] == "WARNING" and issue.get("status") != "ACKNOWLEDGED")
    return {
        "status": "ERROR" if error_count else "WARNING" if warning_count else "PASS",
        "error_count": error_count,
        "warning_count": warning_count,
        "issues": issues,
        "accounts": account_results,
        "summary": {
            "transaction_count": sum(len(bundle["transactions"]) for bundle in accounts.values()),
            "evidence_count": sum(len(bundle["evidence"]) for bundle in accounts.values()),
            "total_income": sum(result["summary"]["total_income"] for result in account_results.values()),
            "total_expense": sum(result["summary"]["total_expense"] for result in account_results.values()),
        },
    }


def generate_combined_validation_report(output_path: Path, validation: dict, accounts: dict[str, dict]) -> None:
    sections: list[str] = []
    for account_id in ("primary", "dues_intake"):
        result = validation["accounts"][account_id]
        summary = result["summary"]
        issue_rows = "\n".join(_validation_issue_html(issue) for issue in result["issues"])
        sections.append(
            f"""
    <section>
      <h2>{html.escape(_account_label(account_id))}</h2>
      <p class="status">상태: {html.escape(result['status'])}</p>
      <table>
        <tr><th>거래 건수</th><td>{summary['transaction_count']:,}건</td></tr>
        <tr><th>이전잔액</th><td>{summary['opening_balance']:,}원</td></tr>
        <tr><th>수입총액</th><td>{summary['total_income']:,}원</td></tr>
        <tr><th>지출총액</th><td>{summary['total_expense']:,}원</td></tr>
        <tr><th>계산 최종잔액</th><td>{summary['computed_closing_balance']:,}원</td></tr>
        <tr><th>장부 최종잔액</th><td>{summary['reported_closing_balance']:,}원</td></tr>
      </table>
      <h3>검증 항목</h3>
      <ul>{issue_rows or '<li>검증 이슈가 없습니다.</li>'}</ul>
    </section>"""
        )
    output_path.write_text(
        f"""<!doctype html>
<html lang="ko">
<head>
  <meta charset="utf-8" />
  <title>ATLAS 통합 검증 리포트</title>
  <style>
    body {{ font-family: Arial, sans-serif; margin: 40px; color: #172026; }}
    h1 {{ margin-bottom: 8px; }} section {{ margin-top: 34px; }}
    table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
    th, td {{ border: 1px solid #c7d0d9; padding: 10px; text-align: left; }}
    th {{ width: 220px; background: #e9f3df; }} .status {{ font-weight: 700; }}
    .ledger-id {{ display: inline-block; margin: 0 6px; padding: 2px 6px; background: #eef2f1; font-weight: 700; }}
  </style>
</head>
<body>
  <h1>ATLAS 동아리연합회 제출 패키지 검증 리포트</h1>
  <p class="status">통합 상태: {html.escape(validation['status'])} · 오류 {validation['error_count']}건 · 경고 {validation['warning_count']}건</p>
  {''.join(sections)}
</body>
</html>
""",
        encoding="utf-8",
    )


def build_combined_submission_package(
    output_root: Path,
    package_id: str,
    club_name: str,
    semester: str,
    treasurer_name: str,
    president_name: str,
    reviewer_name: str | None,
    accounts: dict[str, dict],
    row_capacity: int,
    created_by: str = "system",
) -> dict:
    required_accounts = {"primary", "dues_intake"}
    if set(accounts) != required_accounts:
        raise ValueError("통합 동연 패키지에는 동아리운영계좌(우리은행)와 회비입금계좌(IBK기업은행) 장부가 모두 필요합니다.")

    package_dir = output_root / package_id
    package_dir.mkdir(parents=True, exist_ok=True)
    zip_path = output_root / f"{package_id}.zip"
    ledger_path = package_dir / "동아리 수입지출관리대장.xlsx"
    account_path = package_dir / "동아리 계좌 전체내역.docx"
    evidence_paths = {
        "primary": package_dir / "영수증 및 소명자료 - 동아리운영계좌(우리은행).docx",
        "dues_intake": package_dir / "영수증 및 소명자료 - 회비입금계좌(IBK기업은행).docx",
    }
    report_path = package_dir / "검증_리포트.html"
    manifest_path = package_dir / "manifest.json"

    validation = _combined_validation(accounts)
    ledger_coverage = generate_combined_ledger_workbook(
        ledger_path,
        club_name,
        semester,
        treasurer_name,
        president_name,
        accounts,
        row_capacity,
    )
    evidence_coverage = {
        account_id: generate_evidence_document(
            evidence_paths[account_id],
            club_name,
            treasurer_name,
            reviewer_name,
            accounts[account_id]["transactions"],
            accounts[account_id]["evidence"],
        )
        for account_id in ("primary", "dues_intake")
    }

    combined_evidence: list[Evidence] = []
    for account_id, bundle in accounts.items():
        combined_evidence.extend(bundle["evidence"])
    account_coverage = generate_account_capture_document(
        account_path,
        club_name,
        president_name,
        combined_evidence,
    )
    generate_combined_validation_report(report_path, validation, accounts)

    copied_evidence: list[dict] = []
    copied_sources: list[dict] = []
    account_folders = {"primary": _account_label("primary"), "dues_intake": _account_label("dues_intake")}
    for account_id, bundle in accounts.items():
        evidence_dir = package_dir / "증빙 원본" / account_folders[account_id]
        source_dir = package_dir / "원본자료" / account_folders[account_id]
        evidence_dir.mkdir(parents=True, exist_ok=True)
        source_dir.mkdir(parents=True, exist_ok=True)
        for item in bundle["evidence"]:
            source_path = Path(item.local_path or "")
            if not source_path.exists():
                continue
            target = evidence_dir / f"{item.id}_{normalize_filename(item.filename or source_path.name)}"
            shutil.copy2(source_path, target)
            copied_evidence.append(
                {
                    "id": item.id,
                    "account_id": account_id,
                    "filename": normalize_filename(target.name),
                    "relative_path": normalize_relative_path(target.relative_to(package_dir)),
                    "sha256": file_sha256(target),
                    "transaction_ids": list(item.transaction_ids),
                    "transaction_number": item.transaction_number,
                }
            )
        for source in bundle.get("source_files", []):
            if source.get("kind") != "ledger_workbook":
                continue
            source_path = Path(str(source.get("local_path") or ""))
            if not source_path.exists():
                continue
            target = source_dir / f"{source.get('kind', 'source')}_{normalize_filename(str(source.get('filename') or source_path.name))}"
            shutil.copy2(source_path, target)
            copied_sources.append(
                {
                    "account_id": account_id,
                    "kind": source.get("kind", "source"),
                    "filename": normalize_filename(target.name),
                    "relative_path": normalize_relative_path(target.relative_to(package_dir)),
                    "sha256": file_sha256(target),
                }
            )

    generated_paths = [ledger_path, account_path, *evidence_paths.values(), report_path]
    file_entries = [
        {
            "relative_path": normalize_relative_path(path.relative_to(package_dir)),
            "size": path.stat().st_size,
            "sha256": file_sha256(path),
        }
        for path in generated_paths
    ]
    for item in [*copied_evidence, *copied_sources]:
        path = package_dir / item["relative_path"]
        file_entries.append({"relative_path": item["relative_path"], "size": path.stat().st_size, "sha256": item["sha256"]})

    document_coverage = {
        "ledger": ledger_coverage,
        "evidence_documents": evidence_coverage,
        "account_document": account_coverage,
        "copied_evidence_files": len(copied_evidence),
        "copied_source_files": len(copied_sources),
    }
    manifest = {
        "package_id": package_id,
        "club_name": club_name,
        "semester": semester,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "generated_by": created_by,
        "validation_rule_version": VALIDATION_RULE_VERSION,
        "template_version": TEMPLATE_VERSION,
        "accounts": {
            account_id: {
                "label": _account_label(account_id),
                "snapshot_id": bundle.get("snapshot_id"),
                "ledger_data_hash": bundle.get("ledger_data_hash"),
                "transaction_count": len(bundle["transactions"]),
                "evidence_count": len(bundle["evidence"]),
            }
            for account_id, bundle in accounts.items()
        },
        "document_coverage": document_coverage,
        "validation": validation,
        "evidence": copied_evidence,
        "source_files": copied_sources,
        "files": file_entries,
    }
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in package_dir.rglob("*"):
            if path.is_file():
                archive.write(path, normalize_relative_path(path.relative_to(package_dir)))

    zip_hash = file_sha256(zip_path)
    manifest_hash = file_sha256(manifest_path)
    effective_capacity = int(ledger_coverage["accounts"]["primary"]["row_capacity"])
    return {
        "zip_path": str(zip_path),
        "zip_sha256": zip_hash,
        "manifest_sha256": manifest_hash,
        "manifest": manifest,
        "document_coverage": document_coverage,
        "row_capacity": effective_capacity,
        "validation": validation,
        "artifacts": [
            {"kind": "combined_ledger", "path": str(ledger_path), "sha256": file_sha256(ledger_path)},
            {"kind": "combined_account_history", "path": str(account_path), "sha256": file_sha256(account_path)},
            {"kind": "primary_evidence", "path": str(evidence_paths["primary"]), "sha256": file_sha256(evidence_paths["primary"])},
            {"kind": "dues_evidence", "path": str(evidence_paths["dues_intake"]), "sha256": file_sha256(evidence_paths["dues_intake"])},
            {"kind": "validation_report", "path": str(report_path), "sha256": file_sha256(report_path)},
            {"kind": "manifest", "path": str(manifest_path), "sha256": manifest_hash},
            {"kind": "zip", "path": str(zip_path), "sha256": zip_hash},
        ],
    }


def build_monthly_report_payload(
    report_id: str,
    share_id: str,
    club_name: str,
    month: str,
    opening_balance: int,
    transactions: list[Transaction],
    visible_notes: bool,
    snapshot_id: str | None = None,
    expires_at: str | None = None,
    allow_download: bool = False,
) -> dict:
    summary = public_monthly_summary(opening_balance, transactions)
    public_transactions = []
    for item in transactions:
        public_transactions.append(
            {
                "number": item.number,
                "date": item.date,
                "description": item.description,
                "category": item.category,
                "income": item.income,
                "expense": item.expense,
                "balance": item.balance,
                "note": item.note if visible_notes else "",
            }
        )
    return {
        "id": report_id,
        "share_id": share_id,
        "snapshot_id": snapshot_id,
        "club_name": club_name,
        "month": month,
        "summary": summary,
        "transactions": public_transactions,
        "status": "active",
        "expires_at": expires_at,
        "allow_download": allow_download,
        "content_hash": canonical_sha256({"summary": summary, "transactions": public_transactions}),
    }
