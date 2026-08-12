from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from PIL import Image

from app.services.accounting import CARD_EVIDENCE_METHODS, Evidence, Transaction, public_monthly_summary, validate_ledger
from app.services.documents import build_combined_submission_package, build_submission_package, generate_evidence_document


class AccountingValidationTest(unittest.TestCase):
    def test_valid_ledger_passes_balance_and_evidence_checks(self) -> None:
        transactions = [
            Transaction(1, "2026-03-01", "회비", 100_000, 0, 1_100_000),
            Transaction(2, "2026-03-02", "현수막", 0, 20_000, 1_080_000, evidence_id="ev1", processing_method="카드결제"),
        ]
        evidence = [Evidence("ev1", 2, "receipt.png", "receipt")]

        result = validate_ledger(1_000_000, transactions, evidence, 1_080_000)

        self.assertEqual(result["status"], "PASS")
        self.assertEqual(result["error_count"], 0)
        self.assertEqual(result["summary"]["total_income"], 100_000)
        self.assertEqual(result["summary"]["total_expense"], 20_000)

    def test_invalid_ledger_reports_missing_evidence_and_balance_mismatch(self) -> None:
        transactions = [
            Transaction(1, "2026-03-01", "비품", 0, 30_000, 980_000, processing_method="카드결제"),
        ]

        result = validate_ledger(1_000_000, transactions, [], 960_000)
        codes = {issue["code"] for issue in result["issues"]}

        self.assertEqual(result["status"], "ERROR")
        self.assertIn("BALANCE_CONTINUITY", codes)
        self.assertIn("MISSING_CARD_EVIDENCE", codes)
        self.assertIn("CLOSING_BALANCE_MISMATCH", codes)

    def test_card_payment_and_cancellation_require_receipt_or_explanation(self) -> None:
        transactions = [
            Transaction(1, "2026-03-01", "비품 결제", 0, 30_000, 970_000, processing_method="카드결제"),
            Transaction(2, "2026-03-02", "비품 결제 취소", 30_000, 0, 1_000_000, processing_method="카드결제취소"),
        ]

        result = validate_ledger(1_000_000, transactions, [], 1_000_000)
        missing = [issue for issue in result["issues"] if issue["code"] == "MISSING_CARD_EVIDENCE"]

        self.assertEqual(result["status"], "ERROR")
        self.assertEqual([issue["transaction_number"] for issue in missing], [1, 2])
        self.assertTrue(all("매칭되는 소명자료 또는 영수증이 존재하지 않습니다" in issue["message"] for issue in missing))

    def test_bank_transfer_evidence_is_optional(self) -> None:
        transactions = [
            Transaction(1, "2026-03-01", "대관료", 0, 30_000, 970_000, processing_method="계좌이체"),
        ]

        result = validate_ledger(1_000_000, transactions, [], 970_000)

        self.assertEqual(result["status"], "PASS")
        self.assertNotIn("MISSING_CARD_EVIDENCE", {issue["code"] for issue in result["issues"]})

    def test_evidence_with_same_number_from_other_account_does_not_match(self) -> None:
        transactions = [
            Transaction(
                1,
                "2026-03-01",
                "Card purchase",
                0,
                30_000,
                970_000,
                processing_method=min(CARD_EVIDENCE_METHODS, key=len),
                account_id="primary",
            ),
        ]
        evidence = [Evidence("dues-ev1", 1, "*1* receipt.pdf", "receipt", account_id="dues_intake")]

        result = validate_ledger(1_000_000, transactions, evidence, 970_000)

        self.assertIn("MISSING_CARD_EVIDENCE", {issue["code"] for issue in result["issues"]})

    def test_monthly_summary_hides_private_metadata(self) -> None:
        transactions = [
            Transaction(1, "2026-03-01", "회비", 200_000, 0, 1_200_000),
            Transaction(2, "2026-03-10", "간식", 0, 50_000, 1_150_000, note="내부 메모"),
        ]

        summary = public_monthly_summary(1_000_000, transactions)

        self.assertEqual(summary["closing_balance"], 1_150_000)
        self.assertEqual(len(summary["categories"]["income"]), 1)
        self.assertEqual(len(summary["categories"]["expense"]), 1)
        self.assertNotIn("note", summary["categories"]["expense"][0])
        self.assertNotIn("evidence_id", summary["categories"]["expense"][0])


class PackageGenerationTest(unittest.TestCase):
    def test_evidence_document_marks_cancelled_payments_as_negative_expense(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            card_receipt = root / "#1# card-cancel.png"
            interest_receipt = root / "#2# interest-cancel.png"
            for path in (card_receipt, interest_receipt):
                Image.new("RGB", (640, 420), "white").save(path)
            transactions = [
                Transaction(1, "2026-03-01", "카드 승인 취소", 10_000, 0, 100_000, processing_method="카드결제취소"),
                Transaction(2, "2026-03-02", "이자 입금 취소", 500, 0, 99_500, processing_method="이자입금취소"),
            ]
            evidence = [
                Evidence("card-cancel", 1, card_receipt.name, "receipt", local_path=str(card_receipt)),
                Evidence("interest-cancel", 2, interest_receipt.name, "explanation", local_path=str(interest_receipt)),
            ]
            output_path = root / "evidence.docx"

            generate_evidence_document(output_path, "Aegis", "회계", "검토", transactions, evidence)
            document = Document(output_path)

        self.assertEqual(document.tables[0].cell(4, 2).text, "-10,000원")
        self.assertEqual(document.tables[0].cell(9, 2).text, "-500원")

    def test_combined_package_separates_same_number_transactions_by_account(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            primary_receipt = root / "#1# primary receipt.png"
            dues_receipt = root / "*1* dues receipt.png"
            primary_history = root / "우리은행 거래내역.png"
            dues_history = root / "기업은행 거래내역.png"
            for path in (primary_receipt, dues_receipt, primary_history, dues_history):
                Image.new("RGB", (640, 420), "white").save(path)

            accounts = {
                "primary": {
                    "snapshot_id": "snap-primary",
                    "ledger_data_hash": "a" * 64,
                    "opening_balance": 100_000,
                    "expected_closing_balance": 90_000,
                    "period_start": "2026-01-01",
                    "period_end": "2026-06-30",
                    "transactions": [Transaction(1, "2026-03-01", "운영계좌 카드결제", 0, 10_000, 90_000, processing_method="카드결제", account_id="primary")],
                    "evidence": [
                        Evidence("primary-receipt", 1, primary_receipt.name, "receipt", account_id="primary", local_path=str(primary_receipt)),
                        Evidence("primary-history", None, primary_history.name, "account_capture", account_id="primary", local_path=str(primary_history)),
                    ],
                    "bank_transactions": [{"occurred_at": "2026-03-01T12:00:00", "description": "Primary bank row", "amount": -10_000, "balance": 90_000, "transaction_type": "카드"}],
                },
                "dues_intake": {
                    "snapshot_id": "snap-dues",
                    "ledger_data_hash": "b" * 64,
                    "opening_balance": 200_000,
                    "expected_closing_balance": 195_000,
                    "period_start": "2026-01-01",
                    "period_end": "2026-06-30",
                    "transactions": [Transaction(1, "2026-03-02", "회비계좌 카드결제", 0, 5_000, 195_000, processing_method="카드결제", account_id="dues_intake")],
                    "evidence": [
                        Evidence("dues-receipt", 1, dues_receipt.name, "receipt", account_id="dues_intake", local_path=str(dues_receipt)),
                        Evidence("dues-history", None, dues_history.name, "account_capture", account_id="dues_intake", local_path=str(dues_history)),
                    ],
                    "bank_transactions": [{"occurred_at": "2026-03-02T12:00:00", "description": "Dues bank row", "amount": -5_000, "balance": 195_000, "transaction_type": "이체"}],
                },
            }

            result = build_combined_submission_package(
                root,
                "pkg-combined",
                "Aegis",
                "2026년 1학기",
                "회계",
                "회장",
                "검토",
                accounts,
                40,
            )
            package_dir = root / "pkg-combined"
            workbook = load_workbook(package_dir / "동아리 수입지출관리대장.xlsx", data_only=False)
            primary_document = Document(package_dir / "영수증 및 소명자료 - 동아리운영계좌(우리은행).docx")
            dues_document = Document(package_dir / "영수증 및 소명자료 - 회비입금계좌(IBK기업은행).docx")
            account_document = Document(package_dir / "동아리 계좌 전체내역.docx")
            with zipfile.ZipFile(result["zip_path"]) as archive:
                names = set(archive.namelist())

        self.assertEqual(workbook.sheetnames, ["동아리운영계좌(우리은행)", "회비입금계좌(IBK기업은행)"])
        self.assertEqual(workbook["동아리운영계좌(우리은행)"]["G10"].value, "운영계좌 카드결제")
        self.assertEqual(workbook["회비입금계좌(IBK기업은행)"]["G10"].value, "회비계좌 카드결제")
        self.assertIn("운영계좌 카드결제", "\n".join(cell.text for table in primary_document.tables for row in table.rows for cell in row.cells))
        self.assertNotIn("회비계좌 카드결제", "\n".join(cell.text for table in primary_document.tables for row in table.rows for cell in row.cells))
        self.assertIn("회비계좌 카드결제", "\n".join(cell.text for table in dues_document.tables for row in table.rows for cell in row.cells))
        self.assertGreaterEqual(len(account_document.inline_shapes), 4)
        self.assertEqual(len(account_document.tables[0].cell(0, 0).paragraphs), 1)
        self.assertEqual(len(account_document.tables[0].cell(0, 1).paragraphs), 1)
        self.assertEqual(result["document_coverage"]["account_document"]["generated_bank_pages"], 2)
        self.assertEqual(result["validation"]["status"], "PASS")
        self.assertIn("동아리 계좌 전체내역.docx", names)
        self.assertIn("동아리 수입지출관리대장.xlsx", names)
        self.assertIn("영수증 및 소명자료 - 동아리운영계좌(우리은행).docx", names)
        self.assertIn("영수증 및 소명자료 - 회비입금계좌(IBK기업은행).docx", names)

    def test_submission_report_marks_missing_card_evidence_as_error(self) -> None:
        transactions = [
            Transaction(1, "2026-03-01", "행사용 비품", 0, 30_000, 970_000, processing_method=min(CARD_EVIDENCE_METHODS, key=len)),
        ]
        with tempfile.TemporaryDirectory() as tmp:
            result = build_submission_package(
                Path(tmp),
                "pkg_missing_card_evidence",
                "Aegis",
                "2026년 1학기",
                "회계",
                "회장",
                "검토",
                1_000_000,
                970_000,
                transactions,
                [],
                40,
            )
            report = (Path(tmp) / "pkg_missing_card_evidence" / "검증_리포트.html").read_text(encoding="utf-8")

        self.assertEqual(result["validation"]["status"], "ERROR")
        self.assertEqual(result["document_coverage"]["evidence_document"]["missing_required_card_numbers"], [1])
        self.assertIn("MISSING_CARD_EVIDENCE", report)
        self.assertIn("장부 1번 카드결제 거래에 매칭되는 소명자료 또는 영수증이 존재하지 않습니다.", report)

    def test_official_ledger_template_rejects_more_than_120_rows(self) -> None:
        transactions = [Transaction(index, "2026-03-01", f"거래 {index}", 0, 1, -index) for index in range(1, 122)]
        with tempfile.TemporaryDirectory() as tmp, self.assertRaisesRegex(ValueError, "최대 120칸"):
            build_submission_package(
                Path(tmp),
                "pkg_too_many",
                "Aegis",
                "2026년 1학기",
                "회계",
                "회장",
                "검토",
                0,
                -121,
                transactions,
                [],
                120,
            )

    def test_submission_package_zip_contains_required_artifacts(self) -> None:
        transactions = [
            Transaction(1, "2026-03-01", "회비", 100_000, 0, 1_100_000),
            Transaction(2, "2026-03-02", "현수막", 0, 20_000, 1_080_000, evidence_id="ev1"),
            Transaction(3, "2026-03-03", "회비 환불 소명", 0, 10_000, 1_070_000, evidence_id="ev2"),
        ]
        evidence = [
            Evidence("ev1", 2, "receipt.png", "receipt"),
            Evidence("ev2", 3, "*3* dues-refund.pdf", "explanation", account_id="dues_intake"),
        ]

        with tempfile.TemporaryDirectory() as tmp:
            result = build_submission_package(
                Path(tmp),
                "pkg_test",
                "Aegis",
                "2026년 1학기",
                "회계",
                "회장",
                "검토",
                1_000_000,
                1_070_000,
                transactions,
                evidence,
                40,
            )

            self.assertEqual(result["validation"]["status"], "PASS")
            with zipfile.ZipFile(result["zip_path"]) as archive:
                names = set(archive.namelist())
            self.assertIn("수입지출관리대장.xlsx", names)
            self.assertIn("영수증_및_소명자료.docx", names)
            self.assertIn("계좌전체내역.docx", names)
            self.assertIn("검증_리포트.html", names)
            self.assertIn("manifest.json", names)
            self.assertEqual(len(result["zip_sha256"]), 64)
            self.assertIn("40칸", result["document_coverage"]["ledger_template"]["template_filename"])
            self.assertIn("영수증 및 소명자료", result["document_coverage"]["evidence_document"]["template_filename"])
            self.assertIn("계좌전체내역", result["document_coverage"]["account_document"]["template_filename"])
            accounts = result["document_coverage"]["evidence_document"]["account_breakdown"]
            self.assertIn("dues_intake", {item["account_id"] for item in accounts})


if __name__ == "__main__":
    unittest.main()
