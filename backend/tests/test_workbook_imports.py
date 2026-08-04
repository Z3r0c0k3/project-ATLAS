from __future__ import annotations

import tempfile
import unicodedata
import unittest
import zipfile
from datetime import datetime
from pathlib import Path

import fitz
from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook, load_workbook
from PIL import Image

import app.main as main
from app.services.accounting import Evidence, Transaction
from app.services.documents import build_submission_package
from app.services.jobs import JobRunner
from app.services.store import JsonStore
from app.services.workbooks import parse_aegis_ledger, parse_toss_bank, reconcile_ledger_bank


def make_ledger(path: Path, transaction_count: int = 2) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "1학기"
    headers = ["NO", "날짜", "내용", "수입", "지출", "잔액", "처리방식", "상세정보"]
    for index, value in enumerate(headers, 2):
        sheet.cell(5, index, value)
    running = 0
    for index in range(transaction_count):
        row = 6 + index
        income = 1_000 if index == 0 else 0
        expense = 200 if index == 1 else 10 if index > 1 else 0
        running += income - expense
        values = [
            index + 1,
            datetime(2026, 3, min(index + 1, 28)),
            "이월금" if index == 0 else f"물품 {index}",
            income or None,
            expense or None,
            running,
            "계좌이체" if index == 0 else "카드결제",
            "결제처: 테스트상점" if expense else "입금자: 전년도",
        ]
        for column, value in enumerate(values, 2):
            sheet.cell(row, column, value)
    workbook.save(path)


def make_bank(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "토스뱅크 거래내역"
    headers = ["거래 일시", "적요", "거래 유형", "거래 기관", "계좌번호", "거래 금액", "거래 후 잔액", "메모"]
    for index, value in enumerate(headers, 2):
        sheet.cell(9, index, value)
    values = [datetime(2026, 3, 2, 12, 0), "테스트상점", "체크카드결제", "", "", -200, 800, "물품 1"]
    for column, value in enumerate(values, 2):
        sheet.cell(10, column, value)
    workbook.save(path)


class WorkbookServiceTest(unittest.TestCase):
    def test_parse_reconcile_and_never_truncate_documents(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            ledger_path = root / "ledger.xlsx"
            bank_path = root / "bank.xlsx"
            receipt_path = root / "2_receipt.png"
            explanation_path = root / "3_explanation.pdf"
            capture_path = root / "account.png"
            make_ledger(ledger_path, transaction_count=41)
            make_bank(bank_path)
            Image.new("RGB", (900, 500), "white").save(receipt_path)
            Image.new("RGB", (900, 1200), "white").save(capture_path)
            pdf = fitz.open()
            for page_number in range(2):
                page = pdf.new_page()
                page.insert_text((72, 72), f"Explanation page {page_number + 1}")
            pdf.save(explanation_path)
            pdf.close()

            ledger = parse_aegis_ledger(ledger_path, period="2026년 1학기")
            bank = parse_toss_bank(bank_path)
            self.assertEqual(ledger["transactions"][1]["processing_method"], "카드결제")
            self.assertEqual(ledger["transactions"][1]["details"], "결제처: 테스트상점")
            reconciliation = reconcile_ledger_bank({**ledger, "transactions": ledger["transactions"][:2]}, bank)
            self.assertEqual(reconciliation["status"], "PASS")
            self.assertEqual(reconciliation["balance_delta"], 0)

            transactions = [Transaction(**item) for item in ledger["transactions"]]
            evidence = [
                Evidence("receipt", 2, unicodedata.normalize("NFD", "2_영수증.png"), "receipt", local_path=str(receipt_path)),
                Evidence("explanation", 3, explanation_path.name, "explanation", local_path=str(explanation_path)),
                Evidence("capture", None, capture_path.name, "account_capture", local_path=str(capture_path)),
            ]
            result = build_submission_package(
                root / "outputs",
                "pkg-test",
                "Aegis",
                "2026년 1학기",
                "회계",
                "회장",
                "검토",
                0,
                ledger["closing_balance"],
                transactions,
                evidence,
                40,
                bank_transactions=bank["transactions"],
                source_files=[
                    {
                        "kind": "ledger_workbook",
                        "filename": unicodedata.normalize("NFD", "원본 장부.xlsx"),
                        "local_path": str(ledger_path),
                    }
                ],
            )
            package_dir = root / "outputs" / "pkg-test"
            output_workbook = load_workbook(package_dir / "수입지출관리대장.xlsx", data_only=False)
            output_sheet = output_workbook.active
            self.assertEqual(result["row_capacity"], 80)
            self.assertEqual(output_sheet["B2"].value, "Aegis 수입지출 관리대장")
            self.assertEqual(output_sheet["B7"].value, "1학기")
            self.assertEqual(output_sheet["G50"].value, "물품 40")
            self.assertEqual(output_sheet["U10"].value, "=T7+K10-Q10")
            self.assertIn("B2:O5", {str(item) for item in output_sheet.merged_cells.ranges})
            self.assertIn("$A$1:$Y$94", str(output_sheet.print_area))
            evidence_document = Document(package_dir / "영수증_및_소명자료.docx")
            account_document = Document(package_dir / "계좌전체내역.docx")
            self.assertEqual(len(evidence_document.tables), 21)
            self.assertIn("번호(수입지출관리대장) : 2", evidence_document.tables[0].cell(0, 3).text)
            self.assertGreaterEqual(len(evidence_document.inline_shapes), 3)
            self.assertEqual(len(account_document.tables), 4)
            self.assertIn("Aegis", {paragraph.text.strip() for paragraph in account_document.paragraphs})
            self.assertGreaterEqual(len(account_document.inline_shapes), 1)
            self.assertEqual(result["document_coverage"]["evidence_document"]["embedded_pages"], 3)
            self.assertIn("80칸", result["document_coverage"]["ledger_template"]["template_filename"])
            self.assertEqual(result["document_coverage"]["copied_evidence_files"], 3)
            self.assertEqual(result["document_coverage"]["copied_source_files"], 1)
            with zipfile.ZipFile(result["zip_path"]) as archive:
                archived_names = archive.namelist()
            self.assertTrue(all(unicodedata.is_normalized("NFC", name) for name in archived_names))
            self.assertIn("증빙자료/receipt_2_영수증.png", archived_names)
            self.assertIn("원본자료/ledger_workbook_원본 장부.xlsx", archived_names)


class WorkbookApiTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.previous_store = main.store
        self.previous_jobs = main.jobs
        self.previous_upload_root = main.UPLOAD_ROOT
        self.previous_output_root = main.OUTPUT_ROOT
        main.store = JsonStore(self.root / "store")
        main.jobs = JobRunner(main.store, workers=1)
        main.UPLOAD_ROOT = self.root / "uploads"
        main.OUTPUT_ROOT = self.root / "outputs"
        self.client = TestClient(main.app)
        response = self.client.post("/auth/login", json={"username": "admin", "role": "admin"})
        self.headers = {"X-ATLAS-Token": response.json()["token"]}

    def tearDown(self) -> None:
        main.jobs.executor.shutdown(wait=True)
        main.store = self.previous_store
        main.jobs = self.previous_jobs
        main.UPLOAD_ROOT = self.previous_upload_root
        main.OUTPUT_ROOT = self.previous_output_root
        self.temp.cleanup()

    def test_upload_import_and_attach_evidence_snapshot(self) -> None:
        ledger_path = self.root / "Aegis-ledger.xlsx"
        bank_path = self.root / "toss.xlsx"
        receipt_path = self.root / "2_receipt.png"
        make_ledger(ledger_path)
        make_bank(bank_path)
        Image.new("RGB", (400, 300), "white").save(receipt_path)
        nfd_ledger_filename = unicodedata.normalize("NFD", "Aegis 회계장부.xlsx")
        nfd_evidence_filename = unicodedata.normalize("NFD", "#2# 영수증.png")

        with ledger_path.open("rb") as stream:
            ledger_upload = self.client.post(
                "/imports/upload",
                headers=self.headers,
                files={"file": (nfd_ledger_filename, stream, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
            )
        with bank_path.open("rb") as stream:
            bank_upload = self.client.post(
                "/imports/upload",
                headers=self.headers,
                files={"file": (bank_path.name, stream, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
            )
        self.assertEqual(ledger_upload.json()["detected_type"], "aegis_ledger")
        self.assertEqual(ledger_upload.json()["filename"], "Aegis 회계장부.xlsx")
        self.assertTrue(unicodedata.is_normalized("NFC", ledger_upload.json()["filename"]))
        self.assertEqual(bank_upload.json()["detected_type"], "toss_bank")

        imported = self.client.post(
            "/imports/workbook-snapshot",
            headers=self.headers,
            json={
                "ledger_upload_id": ledger_upload.json()["id"],
                "bank_upload_id": bank_upload.json()["id"],
                "period": "2026년 1학기",
            },
        )
        self.assertEqual(imported.status_code, 200)
        self.assertEqual(imported.json()["transaction_count"], 2)
        self.assertEqual(imported.json()["reconciliation"]["status"], "PASS")

        with receipt_path.open("rb") as stream:
            evidence = self.client.post(
                "/evidence/upload",
                headers=self.headers,
                data={"kind": "receipt", "amount": "200", "evidence_date": "2026-03-02"},
                files={"file": (nfd_evidence_filename, stream, "image/png")},
            )
        self.assertEqual(evidence.json()["filename"], "#2# 영수증.png")
        self.assertEqual(evidence.json()["transaction_number"], 2)
        self.assertTrue(unicodedata.is_normalized("NFC", evidence.json()["filename"]))
        listed_evidence = self.client.get("/evidence", headers=self.headers).json()
        self.assertEqual(listed_evidence[0]["filename"], "#2# 영수증.png")
        attached = self.client.post(
            f"/ledger-snapshots/{imported.json()['id']}/evidence",
            headers=self.headers,
            json={"evidence_ids": [evidence.json()["id"]]},
        )
        self.assertEqual(attached.status_code, 200)
        self.assertEqual(len(attached.json()["evidence"]), 1)
        self.assertIn(evidence.json()["id"], attached.json()["transactions"][1]["evidence_ids"])


if __name__ == "__main__":
    unittest.main()
